import io

import docx
from pypdf import PdfReader

from app.utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception:
        logger.exception("Could not extract PDF text")
        raise


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts).strip()
    except Exception:
        logger.exception("Could not extract DOCX text")
        raise


def extract_text(file_bytes: bytes, filename: str) -> str:
    suffix = filename.lower().rsplit(".", maxsplit=1)[-1] if "." in filename else ""
    if suffix == "pdf":
        return extract_text_from_pdf(file_bytes)
    if suffix == "docx":
        return extract_text_from_docx(file_bytes)
    raise ValueError("Unsupported document type")

